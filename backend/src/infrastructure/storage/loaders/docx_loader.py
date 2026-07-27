from pathlib import Path
from typing import List
import logging
import docx
from backend.src.infrastructure.storage.loaders.base_loader import BaseDocumentLoader, ExtractedChunk

logger = logging.getLogger("docx_loader")


class DocxDocumentLoader(BaseDocumentLoader):
    """Loader strategy for extracting text from DOCX files using python-docx."""

    def load(self, file_path: Path) -> List[ExtractedChunk]:
        try:
            doc = docx.Document(str(file_path))
            elements = []

            # 1. Paragraph text
            for p in doc.paragraphs:
                txt = p.text.strip()
                if txt:
                    elements.append(txt)

            # 2. Table text
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        elements.append(" | ".join(row_cells))

            full_text = "\n\n".join(elements)
            if not full_text.strip():
                logger.warning(f"No readable text extracted from DOCX file '{file_path}'.")
                return []

            return [
                ExtractedChunk(
                    content=full_text,
                    page_number=1,
                    metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
                )
            ]
        except Exception as e:
            logger.error(f"Failed to parse DOCX file '{file_path}': {str(e)}")
            raise ValueError(f"Could not read DOCX document content: {str(e)}") from e

