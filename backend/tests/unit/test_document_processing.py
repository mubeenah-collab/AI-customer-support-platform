from pathlib import Path
from backend.src.ai.rag.text_splitter import DocumentTextSplitter, clean_text
from backend.src.infrastructure.storage.loaders.document_loader_factory import DocumentLoaderFactory
from backend.src.infrastructure.storage.loaders.text_loader import TextDocumentLoader


def test_clean_text_function():
    dirty = "Hello \x00 World!\n\n\n\nThis  is   a  test."
    cleaned = clean_text(dirty)
    assert "Hello World!" in cleaned
    assert "  " not in cleaned
    assert "\n\n\n" not in cleaned


def test_document_text_splitter():
    splitter = DocumentTextSplitter(chunk_size=50, chunk_overlap=10)
    long_text = "Paragraph 1 with some text.\n\nParagraph 2 with more information for testing chunking behavior."
    chunks = splitter.split_text(long_text, page_number=1)

    assert len(chunks) >= 2
    assert chunks[0].chunk_index == 0
    assert chunks[0].page_number == 1
    assert chunks[0].token_count > 0


def test_text_document_loader(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Company Knowledge Base Content Line 1.\nLine 2.", encoding="utf-8")

    loader = TextDocumentLoader()
    extracted = loader.load(file_path)

    assert len(extracted) == 1
    assert "Company Knowledge Base" in extracted[0].content


def test_document_loader_factory_txt(tmp_path: Path):
    file_path = tmp_path / "manual.txt"
    file_path.write_text("Manual Content", encoding="utf-8")

    extracted = DocumentLoaderFactory.load_document(file_path)
    assert len(extracted) == 1
    assert extracted[0].content == "Manual Content"


def test_docx_document_loader(tmp_path: Path):
    import docx
    file_path = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("DOCX Paragraph Content for Testing")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    doc.save(str(file_path))

    extracted = DocumentLoaderFactory.load_document(file_path)
    assert len(extracted) == 1
    assert "DOCX Paragraph Content for Testing" in extracted[0].content
    assert "Cell A | Cell B" in extracted[0].content

